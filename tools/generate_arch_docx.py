#!/usr/bin/env python3
"""生成 BlockCut 迁移架构图 DOCX 文档"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ── 全局样式 ──
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Microsoft YaHei'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    hs.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

def add_code_block(doc, code_text, font_size=Pt(8)):
    """添加代码块"""
    for i, line in enumerate(code_text.strip().split('\n')):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.1
        if i == 0:
            p.paragraph_format.space_before = Pt(6)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = font_size
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add spacing after code block
    doc.add_paragraph().paragraph_format.space_before = Pt(0)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(2)

def add_styled_table(doc, headers, rows, col_widths=None):
    """添加带样式的表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9)

    # Data rows
    for r, row_data in enumerate(rows):
        row_cells = table.rows[r + 1].cells
        for c, val in enumerate(row_data):
            row_cells[c].text = str(val)
            for p in row_cells[c].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # spacing
    return table

# ====================================================================
# 封面
# ====================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('BlockCut V1.7.1 → OmniFrame')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('迁移架构设计文档')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4A, 0x6E, 0x8E)

doc.add_paragraph()
doc.add_paragraph()

info_lines = [
    f'文档版本: 1.0',
    f'生成日期: {datetime.date.today().isoformat()}',
    '源项目: BlockCut V1.7.1 (Qt 5.12 / C++11)',
    '目标架构: OmniFrame (.NET Framework 4.8 / WinForms / C# 7.3+)',
    '硬件平台: APS168x64 PCI 运动控制卡, 16 轴, Basler 相机',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ====================================================================
# 目录占位
# ====================================================================
doc.add_heading('目录', level=1)
p = doc.add_paragraph('（在 Word 中按 Ctrl+A → F9 刷新目录域）')
p.paragraph_format.space_after = Pt(6)
# Insert TOC field
paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar1 = parse_xml(r'<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls('w')))
run._r.append(fldChar1)
run2 = paragraph.add_run()
instrText = parse_xml(r'<w:instrText {} xml:space="preserve"> TOC \o "1-3" \h \z \u </w:instrText>'.format(nsdecls('w')))
run2._r.append(instrText)
run3 = paragraph.add_run()
fldChar2 = parse_xml(r'<w:fldChar {} w:fldCharType="separate"/>'.format(nsdecls('w')))
run3._r.append(fldChar2)
run4 = paragraph.add_run('（在 Word 中右键点击此处 → 更新域 以生成目录）')
run4.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run5 = paragraph.add_run()
fldChar3 = parse_xml(r'<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls('w')))
run5._r.append(fldChar3)

doc.add_page_break()

# ====================================================================
# 1. 项目概述
# ====================================================================
doc.add_heading('一、项目概述', level=1)

doc.add_heading('1.1 源项目：BlockCut V1.7.1', level=2)
add_styled_table(doc,
    ['属性', '值'],
    [
        ['技术栈', 'Qt 5.12 / C++11 / MSVC 2015 64-bit'],
        ['功能', '精密玻璃/晶圆切割、摆放、点胶、UV 固化上位机'],
        ['硬件', 'APS PCI 运动控制卡 (APS168x64)，16 轴，Basler 相机，25 DO + 36 DI'],
        ['视觉算法', 'Halcon 边缘检测 + 线拟合 (FitLineAlg)'],
        ['通信协议', 'MQTT (Qt5Mqtt) + HTTP MES + AES-128-ECB 加密'],
        ['代码量', '37 .cpp / 26 .h / 20 .ui / ~15,000 行核心逻辑'],
    ],
    col_widths=[4, 12]
)

doc.add_heading('1.2 目标架构：OmniFrame', level=2)
add_styled_table(doc,
    ['属性', '值'],
    [
        ['技术栈', '.NET Framework 4.8 / WinForms / C# 7.3+'],
        ['DI 容器', 'Microsoft.Extensions.DependencyInjection'],
        ['核心模式', 'ISystemManager 门面 + 34 Manager + StationBase 工站'],
        ['运动控制', 'Motion 抽象基类 → GTS / DMC / Inovance / PCIeM60 实现'],
        ['IO 控制', 'IoCtrl 抽象基类'],
        ['视觉接口', 'IVisionSystem (Connect / Disconnect / Detect)'],
        ['配置系统', 'XML (ConfigManager) + AES 加密'],
        ['窗体嵌入', 'MainForm.ShowCachedForm<T>() → TopLevel=false, Dock=Fill'],
    ],
    col_widths=[4, 12]
)

doc.add_page_break()

# ====================================================================
# 2. 系统架构图
# ====================================================================
doc.add_heading('二、系统架构总览', level=1)

doc.add_paragraph(
    '下图为 BlockCut 子系统在 OmniFrame 架构中的位置及内部组件关系。'
)

# ASCII 架构图
arch_diagram = r'''
┌──────────────────────────────────────────────────────────────────┐
│                        MainForm (主窗体)                          │
│  ┌─────────────────────────────────────────────────────────────┐ │
│  │              panel_Main (Dock=Fill 容器)                     │ │
│  │  ┌─────────────────────────────────────────────────────────┐ │ │
│  │  │          BlockCutMainForm (TopLevel=false, Dock=Fill)    │ │ │
│  │  │  ┌──────────────────┬────────────────┐                  │ │ │
│  │  │  │   4 路相机预览     │   控制面板       │                  │ │ │
│  │  │  │   (2×2 Grid)     │  ▶ 启动 (F5)   │                  │ │ │
│  │  │  │                  │  ⏸ 暂停 (F1)   │                  │ │ │
│  │  │  │                  │  ⏹ 停止 (F6)   │                  │ │ │
│  │  │  │                  │  ↺ 复位 (F3)   │                  │ │ │
│  │  │  └──────────────────┴────────────────┘                  │ │ │
│  │  └─────────────────────────────────────────────────────────┘ │ │
│  └─────────────────────────────────────────────────────────────┘ │
│  ┌─────────────┐  ┌─────────────┐  ┌─────────────┐               │
│  │ treeView     │  │ StatusBar   │  │ RoundButton │               │
│  │ ├ 仪表盘     │  │ 运行/报警LED │  │ 启动/暂停    │               │
│  │ ├ BlockCut ▶ │  │ 用户/角色    │  │ 停止/复位    │               │
│  │ └ ...       │  │ 产量/良率    │  │ 登录/帮助    │               │
│  └─────────────┘  └─────────────┘  └─────────────┘               │
└──────────────────────────────────────────────────────────────────┘

┌──────────────────────────────────────────────────────────────────┐
│                   DI Container (IServiceProvider)                 │
│  ┌─────────────────────────────────────────────────────────────┐ │
│  │  Singleton: ISystemManager, IAlarmManager, IMotionManager,  │ │
│  │             IIoManager, IConfigManager, IMqttManager, ...   │ │
│  │  Singleton: IVisionSystem → BlockCutVision                   │ │
│  │  Singleton: BlockCutMesClient, BarcodeScannerClient         │ │
│  │  Transient: Station_Adjust, Station_Safe, Station_Load, ... │ │
│  │  Transient: BlockCutMainForm (per ShowCachedForm)           │ │
│  └─────────────────────────────────────────────────────────────┘ │
└──────────────────────────────────────────────────────────────────┘

┌──────────────────────────────────────────────────────────────────┐
│                       BlockCut 子系统内部                          │
│                                                                  │
│  ┌──────────┐  ┌──────────┐  ┌──────────┐  ┌──────────┐        │
│  │ Station  │  │ Station  │  │ Station  │  │ Station  │        │
│  │ _Adjust  │  │ _Safe    │  │ _Load    │  │ _Load2   │        │
│  │ (核心)   │  │ (安全)   │  │ (上料1)  │  │ (上料2)  │        │
│  └────┬─────┘  └────┬─────┘  └────┬─────┘  └────┬─────┘        │
│       │             │             │             │               │
│  ┌────┴─────┐  ┌────┴─────┐  ┌────┴─────┐  ┌────┴─────┐        │
│  │ Station  │  │ Station  │  │ BlockCut │  │ BlockCut │        │
│  │ _CasselZ │  │_BottomGet│  │ Vision   │  │ MESClient│        │
│  └──────────┘  └──────────┘  └──────────┘  └──────────┘        │
│       │             │             │             │               │
│  ┌────┴─────────────┴─────────────┴─────────────┴────┐          │
│  │            BlockCutStationBase : StationBase       │          │
│  │  ┌──────────────────────────────────────────────┐ │          │
│  │  │  ManualResetEventSlim 暂停/恢复               │ │          │
│  │  │  气缸控制 + DI 10s 超时等待                   │ │          │
│  │  │  单轴/双轴/多轴 绝对/相对运动 + 到位重试       │ │          │
│  │  │  UVW 3 轴联动旋转 (三角函数)                  │ │          │
│  │  │  EmitMessage / EmitWarning / EmitPrompt       │ │          │
│  │  └──────────────────────────────────────────────┘ │          │
│  └──────────────────────────────────────────────────┘          │
│       │                       │                                │
│  ┌────┴────────┐    ┌─────────┴──────────┐                     │
│  │ Motion_APS  │    │   IoCtrl_APS       │                     │
│  │ : Motion    │    │   : IoCtrl         │                     │
│  │ 16 轴 P/Inv │    │  25 DO + 36 DI     │                     │
│  └─────────────┘    └────────────────────┘                     │
└──────────────────────────────────────────────────────────────────┘
'''
add_code_block(doc, arch_diagram, Pt(6))

doc.add_page_break()

# ====================================================================
# 3. 模块映射表
# ====================================================================
doc.add_heading('三、模块映射表 (Qt → C#)', level=1)

doc.add_heading('3.1 运动控制子系统', level=2)
add_styled_table(doc,
    ['Qt C++ 源文件', '迁移目标 C# 文件', '说明'],
    [
        ['Motion/Axis/APS/ApsAxis.h/.cpp', 'MotionIO/Motion_APS.cs', 'APS 卡 P/Invoke 封装，继承 Motion 基类'],
        ['Motion/IOSignal/IOSignal.h/.cpp', 'MotionIO/IoCtrl_APS.cs', 'APS 卡板载 IO，继承 IoCtrl 基类'],
        ['Motion/Motion.h/.cpp', '复用 MotionManager.cs', 'UI 层运动控制面板'],
        ['Motion/Thread/ThreadParent.h/.cpp', 'BlockCutStationBase.cs', '线程基类 → 工站抽象基类'],
        ['Motion/Thread/ThreadCasselZ.cpp', 'Station_CasselZ.cs', 'CasselZ 工站'],
        ['Motion/Thread/ThreadLoad1.cpp', 'Station_Load.cs', '上料工站'],
        ['Motion/Thread/ThreadLoad2.cpp', 'Station_Load2.cs', '二次上料工站'],
        ['Motion/Thread/ThreadAdjust.cpp', 'Station_Adjust.cs', '核心调整/生产工站 (~3041 行)'],
        ['Motion/Thread/ThreadBottomGetX.cpp', 'Station_BottomGet.cs', '底板取放工站'],
        ['Motion/Thread/ThreadSafe.cpp', 'Station_Safe.cs', '安全监控工站'],
        ['BaseData.h', 'BlockCutConstants.cs + BlockCutEnums.cs', '轴/IO/线程枚举、超时常量'],
    ],
    col_widths=[5.5, 5.5, 5]
)

doc.add_heading('3.2 视觉子系统', level=2)
add_styled_table(doc,
    ['Qt C++ 源文件', '迁移目标 C# 文件', '说明'],
    [
        ['Alg/FitLine/FitLineAlg.h/.cpp', 'BlockCutVision.cs', 'Halcon 线拟合，实现 IVisionSystem'],
        ['Vision/FitLine/FitLineWidget.h/.cpp', 'FitLineControl.cs (Phase 2)', 'WinForms 视觉参数控件'],
        ['Camera/CameraWidget.h/.cpp', '复用 Basler .NET SDK', '相机采集'],
        ['Camera/CameraView.h/.cpp', 'CameraViewControl.cs (Phase 2)', '相机显示控件'],
    ],
    col_widths=[5.5, 5.5, 5]
)

doc.add_heading('3.3 MES / 通信子系统', level=2)
add_styled_table(doc,
    ['Qt C++ 源文件', '迁移目标 C# 文件', '说明'],
    [
        ['Include/qmqttclient.h', '扩展 MqttManager.cs', '添加 AES 加密发布/订阅'],
        ['Mes/HttpToMes.h/.cpp', 'BlockCutMesClient.cs', 'HTTP MES 验证 + MQTT 加密上报'],
        ['qaesencryption.h/.cpp', '扩展 Security.cs', 'Aes128EcbEncrypt / Decrypt'],
        ['Setting/Log/Log.h/.cpp', '复用 OperationLogService', 'Excel 日志写入'],
        ['TcpSocket/', 'BarcodeScannerClient.cs', '扫码枪 TCP 客户端'],
    ],
    col_widths=[5.5, 5.5, 5]
)

doc.add_heading('3.4 UI 子系统', level=2)
add_styled_table(doc,
    ['Qt C++ (Qt UI)', '迁移目标 C# (WinForms)', '说明'],
    [
        ['MainWindow.h/.cpp/.ui (~2228 行)', 'BlockCutMainForm.cs', '主窗体，嵌入 panel_Main'],
        ['Work/Work.h/.cpp/.ui', 'WorkSelectDialog.cs (Phase 2)', '工单选择对话框'],
        ['Statistic/Statistic.h/.cpp/.ui', 'StatisticsControl.cs (Phase 2)', '图表统计面板'],
        ['Setting/ 相关', 'SettingsForm.cs (Phase 2)', '参数设置表单'],
    ],
    col_widths=[5.5, 5.5, 5]
)

doc.add_page_break()

# ====================================================================
# 4. 核心类设计
# ====================================================================
doc.add_heading('四、核心类设计', level=1)

doc.add_heading('4.1 类继承关系', level=2)

class_diagram = r'''
StationBase (OmniFrame.Core)
├── DoExecute() — 模板方法，BlockCut 工站不直接使用
│
└── BlockCutStationBase : StationBase
    ├── Motion_APS Motion       — APS 运动控制卡
    ├── IoCtrl_APS Io           — APS IO 控制
    ├── ManualResetEventSlim    — 暂停/恢复 (替代 busy-wait)
    ├── ManualResetEventSlim    — 用户确认信号
    ├── Pause() / Resume()      — 暂停控制
    ├── CheckPause(token)       — 循环内检查暂停
    ├── EmitMessage(msg)        — 日志/状态事件
    ├── EmitWarning(code,msg)   — 报警事件
    ├── EmitPrompt(msg)         — 用户提示事件
    ├── SetOneCylinderAsync()   — 单气缸 + DI 10s 超时
    ├── SetOneCylinder2Async()  — 双 DI 气缸
    ├── SetCylinderForTwoDoAsync() — 双 DO 气缸
    ├── OneAxisMoveAbsAsync()   — 单轴绝对运动 + 到位重试
    ├── TwoAxisMoveAbsAsync()   — 双轴联动
    ├── MultiAxisMoveAbsAsync() — 多轴联动
    ├── UVWRotateAsync()        — UVW 3 轴三角函数旋转
    ├── DoExecute() => true     — 继承抽象方法桩
    └── abstract RunAsync(CancellationToken)
        │
        ├── Station_Adjust  — 核心生产工站 (~580 行)
        │   ├── BottomCameraGetAngleAsync() — 2 点拍照 + 线方程角度
        │   ├── SliceCameraGetAngleAsync()  — 片源对准迭代
        │   ├── AdjustSliceAsync()          — Y1/Y2 迭代调整 (0.8 系数)
        │   ├── DispensingOnceAsync()       — 点胶 + UV 固化 (3s)
        │   ├── UVWInitAsync()              — UVW 初始角度计算
        │   ├── GoSafeAsync()               — 各轴回安全位
        │   └── events: OnCameraCapture, OnMesValidate, OnSweepCode
        │
        ├── Station_Safe     — 安全监控工站
        │   └── 1ms 轮询 DI_Door + DI_Safe → UV 关 → 报警 → 暂停
        │
        ├── Station_BottomGet — 底板取放工站
        │   └── 送空底板 / 取满板 / 夹紧/松开气缸
        │
        ├── Station_Load     — 上料工站 1 (骨架)
        ├── Station_Load2    — 上料工站 2 (骨架)
        └── Station_CasselZ  — CasselZ 工站 (骨架)

BlockCutVision : IVisionSystem
├── Connect(ip, port)       — Halcon License 初始化
├── Disconnect()            — 释放 Halcon 资源
├── Detect(VisionParams)    — 通用检测 → 委托 FitLine → VisionResult
├── FitLine(image, roi, params) → (Point2D, Point2D, angle)
├── DetectAngle(image, roi) → double
├── GetMaxGray(image) → int
├── GetAvgGray(image) → double
├── AdjustExposure(maxGray, exposureUs) → newExposureUs
├── MeasureHeightPoint(image, roi) → double
├── MeasureHeightGrid(images[], rois[]) → double[9]
└── 参数类型: FitLineParams, ROIRegion
'''
add_code_block(doc, class_diagram, Pt(7))

doc.add_heading('4.2 Motion_APS : Motion (P/Invoke 封装)', level=2)
doc.add_paragraph('17 个 DllImport 方法从 APS168x64.dll 导入，CallingConvention = Cdecl。')
add_styled_table(doc,
    ['DllImport 方法', 'C 原型', '说明'],
    [
        ['APS_init_board', 'int APS_init_board(int board_id)', '初始化板卡'],
        ['APS_close_board', 'int APS_close_board(int board_id)', '关闭板卡'],
        ['APS_set_axis_param', 'int APS_set_axis_param(int, int, double, ...)', '设置轴参数 (加速度/减速度/速度)'],
        ['APS_absolute_move', 'int APS_absolute_move(int, int, double, int)', '绝对运动'],
        ['APS_relative_move', 'int APS_relative_move(int, int, double, int)', '相对运动'],
        ['APS_stop_move', 'int APS_stop_move(int, int, int)', '停止轴'],
        ['APS_emg_stop', 'int APS_emg_stop(int)', '急停'],
        ['APS_get_position', 'int APS_get_position(int, int, ref double)', '获取当前位置'],
        ['APS_get_axis_state', 'int APS_get_axis_state(int, int, ref int)', '获取轴状态'],
        ['APS_servo_on', 'int APS_servo_on(int, int, int)', '伺服使能'],
        ['APS_servo_off', 'int APS_servo_off(int, int, int)', '伺服去使能'],
        ['APS_reset_axis_alarm', 'int APS_reset_axis_alarm(int, int)', '清除轴报警'],
        ['APS_set_soft_limit', 'int APS_set_soft_limit(int, int, double, double)', '设置软限位'],
        ['APS_enable_soft_limit', 'int APS_enable_soft_limit(int, int, int)', '启停软限位'],
        ['APS_get_motion_profile_time', 'int APS_get_motion_profile_time(...)', '计算运动时间'],
        ['APS_move_linear_abs', 'int APS_move_linear_abs(int, int[], double[], int)', '多轴线性绝对运动'],
        ['APS_move_linear_rel', 'int APS_move_linear_rel(int, int[], double[], int)', '多轴线性相对运动'],
    ],
    col_widths=[5, 6, 5]
)

doc.add_heading('4.3 IoCtrl_APS : IoCtrl (板载 IO)', level=2)
doc.add_paragraph(
    '5 个 DllImport + 扩展方法。25 DO 输出点 + 36 DI 输入点，全部映射自 BaseData.h #define 宏。'
)
add_styled_table(doc,
    ['方法', '说明'],
    [
        ['APS_set_do_bit / APS_get_di_bit', '单点 DO 输出 / DI 输入'],
        ['APS_set_do_port / APS_get_di_port', '端口批量读写'],
        ['APS_get_do_port', 'DO 端口状态回读'],
        ['SetDO(int index, bool val)', '带日志的 DO 设置'],
        ['GetDI(int index)', 'DI 读取'],
        ['WaitDIAsync(di, expected, timeout, desc, token)', '等待 DI 到达期望状态 (替代 10s 超时)'],
        ['GreenLight / RedLight / YellowLight / Buzzer', '三色灯 + 蜂鸣器快捷方法'],
    ],
    col_widths=[7, 9]
)

doc.add_heading('4.4 Station_Adjust 主流程图', level=2)
flow = r'''
RunAsync(CancellationToken token) 主循环:

  ┌─────────────────────────────────────────────┐
  │  InitParams()          → 初始化参数         │
  │  GoSafeAsync()         → 各轴回安全位       │
  │  CheckIsBottom()       → 检查底板就绪       │
  │  SendOrGetBottom()     → 送空/取满底板      │
  └──────────────┬──────────────────────────────┘
                 ↓
  ┌─────────────────────────────────────────────┐
  │  while (!token.IsCancellationRequested)     │
  │  ┌───────────────────────────────────────┐  │
  │  │ SweepCode()        → 扫码获取条码    │  │
  │  │ MesValidateAsync() → MES 卡片验证     │  │
  │  │ LoadRecipeParams() → 加载配方参数     │  │
  │  │ UVWInitAsync()     → UVW 初始角度     │  │
  │  └───────────────────────────────────────┘  │
  │                 ↓                           │
  │  for row = 0..rows:                         │
  │    for col = 0..cols:                       │
  │      ┌─────────────────────────────────┐    │
  │      │ CheckHeightRange() → 9点测高    │    │
  │      │ AdjustBottom()     →底板角度调整 │    │
  │      │ SetSliceAsync()    →放片源       │    │
  │      │ BottomCameraGetAngleAsync()       │    │
  │      │   → 2点拍照 + 线方程拟合          │    │
  │      │   → atan(dA/dB) 计算角度          │    │
  │      │ AdjustSliceAsync()                │    │
  │      │   → SliceCameraGetAngleAsync()    │    │
  │      │   → Y1/Y2 迭代调整 (0.8 系数)    │    │
  │      │ DispensingOnceAsync()             │    │
  │      │   → 点胶 + UV 固化 (3s)           │    │
  │      └─────────────────────────────────┘    │
  │  ┌───────────────────────────────────────┐  │
  │  │ UpdateUPH() → 更新产能统计            │  │
  │  │ GoSafeAsync() → 回安全位              │  │
  │  └───────────────────────────────────────┘  │
  └─────────────────────────────────────────────┘
'''
add_code_block(doc, flow, Pt(7))

doc.add_page_break()

# ====================================================================
# 5. 16 轴轴号映射
# ====================================================================
doc.add_heading('五、16 轴轴号映射表', level=1)

add_styled_table(doc,
    ['eAxisID 枚举', '轴号', '名称', '功能说明'],
    [
        ['CasselZ', '0', 'CasselZ', 'Cassel Z 升降'],
        ['LoadX', '1', 'LoadX', '上料 X 轴'],
        ['LoadY', '2', 'LoadY', '上料 Y 轴'],
        ['AdjustX', '3', 'AdjustX', '调整 X 轴'],
        ['AdjustY1', '4', 'AdjustY1', '调整 Y1 轴'],
        ['AdjustY2', '5', 'AdjustY2', '调整 Y2 轴'],
        ['BottomU', '6', 'BottomU', 'UVW 平台 U 轴'],
        ['BottomV', '7', 'BottomV', 'UVW 平台 V 轴'],
        ['BottomW', '8', 'BottomW', 'UVW 平台 W 轴'],
        ['DisX', '9', 'DisX', '点胶 X 轴'],
        ['DisY', '10', 'DisY', '点胶 Y 轴'],
        ['DisZ', '11', 'DisZ', '点胶 Z 轴'],
        ['CameraX', '12', 'CameraX', '相机 X 轴'],
        ['CameraZ', '13', 'CameraZ', '相机 Z 轴'],
        ['BottomY', '14', 'BottomY', '底板 Y 轴'],
        ['BottomGetX', '15', 'BottomGetX', '底板取料 X 轴'],
    ],
    col_widths=[3, 2, 3, 8]
)

doc.add_paragraph()

# ====================================================================
# 6. 关键技术映射
# ====================================================================
doc.add_heading('六、关键技术映射', level=1)

doc.add_heading('6.1 线程模型迁移', level=2)
add_styled_table(doc,
    ['Qt C++ 模式', '.NET / C# 对应', '性能特性'],
    [
        ['QThread + while(!isInterruptionRequested())', 'Task.Run(() => Station.RunAsync(CancellationToken))', '非阻塞 async/await'],
        ['QThread::msleep(n)', 'await Task.Delay(n, token)', '不占用线程池线程'],
        ['QThread::requestInterruption()', 'CancellationTokenSource.Cancel()', '协作式取消'],
        ['CheckPause() — busy spin loop', 'ManualResetEventSlim.Wait(token)', '零 CPU 等待'],
        ['WaitForUser() — busy spin loop', 'ManualResetEventSlim.Wait(token) + UI event', '事件驱动'],
        ['connect(sender, SIGNAL, receiver, SLOT)', 'sender.Event += receiver.Handler', '类型安全编译期检查'],
        ['QMutex + QMutexLocker', 'lock(_mutex) / SemaphoreSlim', '更轻量'],
        ['QMap<QString, QVariant>', 'Dictionary<string, object>', '泛型强类型'],
    ],
    col_widths=[6, 6, 4]
)

doc.add_heading('6.2 信号/槽 → C# Event 映射', level=2)
add_styled_table(doc,
    ['Qt 信号 (MainWindow)', 'C# Event (BlockCutMainForm)', '方向'],
    [
        ['GoOnRunOneThread', 'StationManager.RequestResume', 'UI → 业务'],
        ['EmitAlarmMessage(int, QString)', 'AlarmManager.AlarmOccurred(AlarmInfo)', '业务 → UI'],
        ['NoticeBottomSweepCode(QString)', 'BarcodeScannerClient.OnCodeScanned(string)', '外部 → 业务'],
        ['EmitCameraMessage(QString)', 'CameraManager.OnMessage(string)', '业务 → UI'],
        ['UpdateDisOnceTime(QString)', 'UIManager.OnStatusUpdate(string)', '业务 → UI'],
        ['ChangeWork(bool, QMap, QString, QVector)', 'WorkManager.OnWorkChanged(WorkEventArgs)', '业务 → UI'],
    ],
    col_widths=[6, 6, 4]
)

doc.add_heading('6.3 配置格式迁移 (INI → XML)', level=2)
ini_example = r'''
# INI 原格式
[Work]
workID=W001
machID=M001

[AxisParam]
CasselZ_speed=100.0
LoadX_speed=200.0
'''
add_code_block(doc, ini_example, Pt(8))

xml_example = r'''
<!-- XML 目标格式 -->
<BlockCut>
  <Work workID="W001" machID="M001" />
  <Axes>
    <Axis name="CasselZ" index="0" speed="100.0" acc="500" dec="500" />
    <Axis name="LoadX"   index="1" speed="200.0" acc="500" dec="500" />
    ...
  </Axes>
  <IO>
    <DO index="0" name="JigYCylinderOut" />
    ...
  </IO>
  <Recipes>
    <Recipe code="A001" rows="5" cols="4" plateWidth="300" plateHeight="200" />
    ...
  </Recipes>
</BlockCut>
'''
add_code_block(doc, xml_example, Pt(8))

doc.add_page_break()

# ====================================================================
# 7. 数据流图
# ====================================================================
doc.add_heading('七、数据流图', level=1)

doc.add_heading('7.1 生产数据流', level=2)
data_flow = r'''
┌──────────┐   TCP (CR/LF)   ┌────────────────┐
│  扫码枪   │ ───────────────→│ BarcodeScanner │
└──────────┘                 │ Client         │
                             └───────┬────────┘
                                     │ OnCodeScanned 事件
                                     ↓
┌──────────┐   HTTP GET      ┌────────────────┐
│ MES 服务器 │←──────────────│ BlockCutMes    │←─── Station_Adjust
│           │  validateCard  │ Client         │    MesValidateAsync()
│           │──────────────→│                │
│           │  JSON 响应      │                │
└──────────┘                 └───────┬────────┘
                                     │ MQTT Publish (AES-128-ECB)
                                     ↓
┌──────────┐   MQTT Subscribe ┌────────────────┐
│ MQTT     │←────────────────│ MqttManager    │
│ Broker   │   machine/status │ (IMqttManager) │
│          │   machine/work   │                │
└──────────┘                 └────────────────┘

┌─────────────────────────────────────────────────────────────┐
│                    Station_Adjust 内部数据流                   │
│                                                             │
│  扫码 ─→ MES验证 ─→ 配方加载 ─→ 9点测高 ─→ UVW旋转        │
│    │                                              │         │
│    ↓                                              ↓         │
│  barcode ────→ fileId                      angle ─→ 对齐    │
│                               BlockCutVision.FitLine()      │
│                                     │                       │
│  相机拍照 ─→ 线拟合 ─→ 角度 ─→ Y1/Y2调整 ─→ 点胶 ─→ UV    │
│       ↑                                                    │
│  OnCameraCapture 事件 ─→ BlockCutVision.DetectAngle()       │
└─────────────────────────────────────────────────────────────┘
'''
add_code_block(doc, data_flow, Pt(7))

doc.add_heading('7.2 安全监控数据流', level=2)
safety_flow = r'''
┌──────────────────────────────────────────────────────┐
│                  Station_Safe.RunAsync()              │
│                                                      │
│  1ms 轮询 ──→ DI_Door(31)  ──→ 门吸状态             │
│           ──→ DI_Safe(32)  ──→ 光幕状态             │
│                                                      │
│  不安全 ──→ DO_UV=false    (关闭紫外)               │
│         ──→ EmitWarning()  (报警)                   │
│         ──→ Pause()        (暂停所有工站)            │
│                                                      │
│  恢复   ──→ OnSafetyRestored 事件                    │
│         ──→ ResumeAll()     (恢复所有工站)            │
└──────────────────────────────────────────────────────┘
'''
add_code_block(doc, safety_flow, Pt(7))

doc.add_page_break()

# ====================================================================
# 8. 文件清单
# ====================================================================
doc.add_heading('八、文件清单', level=1)

doc.add_heading('8.1 新增文件 (15 个已创建)', level=2)
add_styled_table(doc,
    ['类别', '文件', '行数', '状态'],
    [
        ['运动控制', 'MotionIO/Motion_APS.cs', '~250', '✅ 完成'],
        ['IO 控制', 'MotionIO/IoCtrl_APS.cs', '~180', '✅ 完成'],
        ['常量/枚举', 'OmniFrame.Core/BlockCut/BlockCutEnums.cs', '~90', '✅ 完成'],
        ['常量/枚举', 'OmniFrame.Core/BlockCut/BlockCutConstants.cs', '~100', '✅ 完成'],
        ['工站基类', 'OmniFrame.Core/BlockCut/BlockCutStationBase.cs', '~450', '✅ 完成'],
        ['生产工站', 'OmniFrame.Core/BlockCut/Station_Adjust.cs', '~900', '✅ 完成'],
        ['安全工站', 'OmniFrame.Core/BlockCut/Station_Safe.cs', '~90', '✅ 完成'],
        ['底板工站', 'OmniFrame.Core/BlockCut/Station_BottomGet.cs', '~120', '✅ 完成'],
        ['CasselZ', 'OmniFrame.Core/BlockCut/Station_CasselZ.cs', '~40', '骨架'],
        ['上料1', 'OmniFrame.Core/BlockCut/Station_Load.cs', '~40', '骨架'],
        ['上料2', 'OmniFrame.Core/BlockCut/Station_Load2.cs', '~40', '骨架'],
        ['视觉', 'OmniFrame.Core/BlockCut/BlockCutVision.cs', '~210', '✅ 完成'],
        ['MES', 'OmniFrame.Core/BlockCut/BlockCutMesClient.cs', '~230', '✅ 完成'],
        ['扫码枪', 'OmniFrame.Core/BlockCut/BarcodeScannerClient.cs', '~120', '✅ 完成'],
        ['共享类型', 'OmniFrame.Core/BlockCut/BlockCutTypes.cs', '~35', '✅ 完成'],
        ['主窗体', 'OmniFrame/BlockCutMainForm.cs', '~500', '✅ 完成'],
    ],
    col_widths=[3, 7, 2, 2.5]
)

doc.add_heading('8.2 修改现有文件 (4 个)', level=2)
add_styled_table(doc,
    ['文件', '变更内容'],
    [
        ['Common/Security.cs', '添加 Aes128EcbEncrypt / Aes128EcbDecrypt (+60 行)'],
        ['OmniFrame.Core/AdvancedFeatures/MqttManager.cs', '添加 Subscribe(Action) / Unsubscribe (+15 行)'],
        ['OmniFrame/Program.cs', '注册 BlockCut 全部 DI 服务 (+30 行)'],
        ['OmniFrame/MainForm.cs', '树节点 + ShowBlockCutForm() (+15 行)'],
    ],
    col_widths=[7, 9]
)

doc.add_page_break()

# ====================================================================
# 9. 实施阶段
# ====================================================================
doc.add_heading('九、实施阶段', level=1)

doc.add_heading('Phase 1 — MVP (约 12 周 / 60 人天)', level=2)
add_styled_table(doc,
    ['序号', '任务', '产出物', '估时'],
    [
        ['1.1', 'APS 卡 P/Invoke 封装', 'Motion_APS.cs, IoCtrl_APS.cs', '5d'],
        ['1.2', 'BlockCut 常量/枚举定义', 'BlockCutConstants.cs, BlockCutEnums.cs', '1d'],
        ['1.3', 'BlockCutStationBase 基类', 'BlockCutStationBase.cs', '3d'],
        ['1.4', 'Station_BottomGet', '底板取放工站', '4d'],
        ['1.5', 'Station_Load', '上料工站', '3d'],
        ['1.6', 'Station_Load2', '二次上料工站', '2d'],
        ['1.7', 'Station_CasselZ', 'CasselZ 工站', '2d'],
        ['1.8', 'Station_Safe', '安全监控工站', '3d'],
        ['1.9', 'Station_Adjust（核心）', '核心生产工站', '10d'],
        ['1.10', 'Halcon 视觉 .NET 封装', 'BlockCutVision.cs', '5d'],
        ['1.11', 'Security AES-128-ECB', '扩展 Security.cs', '1d'],
        ['1.12', 'MQTT MES 扩展', '扩展 MqttManager.cs', '2d'],
        ['1.13', 'MES 客户端', 'BlockCutMesClient.cs', '3d'],
        ['1.14', 'WinForms 主界面', 'BlockCutMainForm.cs', '5d'],
        ['1.15', '配置迁移工具', 'INI → XML', '2d'],
        ['1.16', 'DI 注册 + 系统集成', 'Program.cs, MainForm.cs', '2d'],
        ['1.17', '集成测试 (模拟)', '测试报告', '7d'],
    ],
    col_widths=[1.2, 5.5, 5.5, 1.5]
)

doc.add_heading('Phase 2 — 完整功能 (约 10 周 / 53 人天)', level=2)
add_styled_table(doc,
    ['序号', '任务', '产出物', '估时'],
    [
        ['2.1', '扫码枪 TCP 客户端', 'BarcodeScannerClient.cs', '2d'],
        ['2.2', '统计图表面板', 'StatisticsControl.cs', '3d'],
        ['2.3', '工单选择对话框', 'WorkSelectDialog.cs', '2d'],
        ['2.4', 'FitLine 视觉参数控件', 'FitLineControl.cs', '3d'],
        ['2.5', '参数设置 UI', 'SettingsForm.cs', '3d'],
        ['2.6', '角色权限 UI 切换', '扩展 PermissionManager + UI', '2d'],
        ['2.7', '报警代码映射', 'BlockCutAlarmMap.cs', '1d'],
        ['2.8', '运动控制 UI 面板', '复用 + 定制 MotionControl', '3d'],
        ['2.9', 'Basler 相机 .NET SDK', '扩展 BlockCutVision.cs', '3d'],
        ['2.10', '产量统计 + UPH', 'ProductivityTracker.cs', '2d'],
        ['2.11', '定时器服务', 'BlockCutTimerService.cs', '1d'],
        ['2.12', '日志/Excel 写入', '扩展 OperationLogService', '2d'],
        ['2.13', 'MQTT 订阅回复', '扩展 MqttManager.cs', '2d'],
        ['2.14', '真机调试', '16 轴 + IO + 相机', '10d'],
        ['2.15', '性能优化 + 稳定性', '—', '5d'],
    ],
    col_widths=[1.2, 5.5, 5.5, 1.5]
)

p = doc.add_paragraph()
run = p.add_run('总计：约 22 周 / ~113 人天')
run.font.bold = True
run.font.size = Pt(11)

doc.add_page_break()

# ====================================================================
# 10. 风险与对策
# ====================================================================
doc.add_heading('十、风险点与对策', level=1)
add_styled_table(doc,
    ['风险', '等级', '对策'],
    [
        ['APS SDK P/Invoke 复杂，参数不正确', '高', '先编写 P/Invoke 单元测试，逐方法验证参数签名'],
        ['Halcon .NET 版本差异（原 Halcon 10）', '高', '先验证 Halcon .NET SDK 兼容性，确认 License 可用'],
        ['ThreadAdjust::Run() 3041 行业务逻辑复杂', '高', '拆分为 15+ 私有方法，逐方法迁移 + 单元测试'],
        ['INI 配置项散落在多文件中', '中', '先做 INI 全量收集整理，再统一转为 XML'],
        ['气缸 10s 超时逻辑在业务代码中耦合', '中', '在 BlockCutStationBase 中统一封装超时处理'],
        ['扫码枪 TCP 多端口监听', '中', '使用独立的 BarcodeScannerClient + Task'],
        ['Basler Pylon SDK .NET 版本适配', '低', 'Pylon 5.1+ 有官方 .NET wrapper，直接引用'],
        ['Excel 日志格式兼容', '低', '保持与 CLog::Instance() 相同列结构'],
    ],
    col_widths=[6, 2, 8]
)

doc.add_page_break()

# ====================================================================
# 11. 前置验证清单
# ====================================================================
doc.add_heading('十一、前置验证清单', level=1)
doc.add_paragraph('在正式开始 Phase 1 编码前，建议先完成以下验证：')

checklist = [
    ('[ ]', 'Halcon .NET SDK 安装 + License 验证（原项目使用 Halcon 10，确认 .NET 版可用）'),
    ('[ ]', 'APS SDK (APS168x64.dll) 存在性 + C# 加载测试'),
    ('[ ]', 'Basler Pylon .NET SDK 安装 + 相机取图测试'),
    ('[ ]', '现有设备 API 手册准备就绪（轴参数、IO 地址映射）'),
    ('[ ]', 'MES 接口文档 / URL / AES Key 确认'),
    ('[ ]', '生产环境 D:/list.ini 内容备份'),
]
for chk, item in checklist:
    p = doc.add_paragraph()
    run = p.add_run(f'{chk} {item}')
    run.font.size = Pt(10)

# ====================================================================
# 页脚
# ====================================================================
doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— 文档结束 —')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(9)

# ── 保存 ──
output_path = '/Users/james/Desktop/Upper/Upper/Upper/Upper/docs/architecture/BlockCut_Architecture_Design.docx'
doc.save(output_path)
print(f'DOCX 已生成: {output_path}')
